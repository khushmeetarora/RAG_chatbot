"""
Final Year Bachelor's Project: Chat with Multiple Audio & Documents

Author: Khushmeey
Date: 2024-04-15
Course: Bachelor of Technology in Computer Science Engineering

Description:
This Streamlit application allows users to upload multiple audio (MP3, WAV, etc.)
and document (PDF, DOCX, TXT) files. It processes these files by:
1. Transcribing audio using OpenAI's Whisper model (via Transformers library).
2. Extracting text from documents using PyPDFLoader, python-docx, and textract.
3. Generating embeddings for the text content using Sentence Transformers.
4. Storing these embeddings in individual FAISS vector stores, cached locally.
5. Merging the vector stores of all currently active files into a single context.
6. Utilizing a Large Language Model (LlamaCpp with a Mistral-based model)
   and LangChain's ConversationalRetrievalChain to answer user questions based
   on the combined content of the uploaded files.

Key Technologies:
- Streamlit (Web Framework)
- Transformers (Whisper ASR)
- LlamaCpp (Local LLM Inference)
- Sentence Transformers (Embeddings)
- FAISS (Vector Store)
- LangChain (Orchestration, RAG)
- Streamlit-Chat (UI Component)
"""

import streamlit as st
import torch
from transformers import AutoModelForSpeechSeq2Seq, AutoProcessor, pipeline
from langchain.chains import ConversationalRetrievalChain
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.memory import ConversationBufferMemory
from langchain.vectorstores import FAISS
from langchain.llms import LlamaCpp
from langchain.schema import Document
from langchain.document_loaders import PyPDFLoader
from docx import Document as DocxDocument # Rename to avoid conflict
import tempfile
import textract
import os
import pickle
import hashlib
from streamlit_chat import message
import logging
from typing import List, Dict, Optional, Tuple

# Import configuration variables
import config

# --- Setup ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
# Ensure cache directory exists
os.makedirs(config.VECTOR_STORE_CACHE_DIR, exist_ok=True)

# --- Model Loading (Cached) ---

@st.cache_resource
def load_whisper_pipeline():
    """
    Loads the Whisper Automatic Speech Recognition (ASR) pipeline.

    Uses Hugging Face Transformers for loading the specified model and processor.
    Caches the pipeline resource using Streamlit for efficiency across reruns.
    Handles device selection (CUDA if available, else CPU).

    Returns:
        Optional[transformers.Pipeline]: The initialized ASR pipeline object,
                                         or None if loading fails.

    Raises:
        Logs errors and displays Streamlit error messages on failure.
    """
    try:
        device = "cuda" if torch.cuda.is_available() else "cpu"
        torch_dtype = torch.float16 if torch.cuda.is_available() else torch.float32
        logging.info(f"Loading Whisper model ({config.WHISPER_MODEL_ID}) on device: {device}")

        model = AutoModelForSpeechSeq2Seq.from_pretrained(
            config.WHISPER_MODEL_ID,
            torch_dtype=torch_dtype,
            low_cpu_mem_usage=True,
            # use_safetensors=True # Consider enabling if model supports it
        )
        model.to(device)
        processor = AutoProcessor.from_pretrained(config.WHISPER_MODEL_ID)

        pipe = pipeline(
            "automatic-speech-recognition",
            model=model,
            tokenizer=processor.tokenizer,
            feature_extractor=processor.feature_extractor,
            max_new_tokens=config.WHISPER_MAX_NEW_TOKENS,
            chunk_length_s=config.WHISPER_CHUNK_LENGTH_S,
            batch_size=config.WHISPER_BATCH_SIZE,
            return_timestamps=config.WHISPER_RETURN_TIMESTAMPS,
            torch_dtype=torch_dtype,
            device=device,
        )
        logging.info("Whisper ASR pipeline loaded successfully.")
        return pipe
    except Exception as e:
        logging.error(f"Error loading Whisper model ({config.WHISPER_MODEL_ID}): {e}", exc_info=True)
        st.error(f"Error loading Whisper ASR model: {e}. Check model ID, network connection, and available memory.")
        return None

@st.cache_resource
def load_llama_cpp():
    """
    Loads the LlamaCpp language model from the specified GGUF file path.

    Uses Streamlit's caching to prevent reloading on every script run.
    Checks if the model file exists before attempting to load.

    Returns:
        Optional[LlamaCpp]: An initialized LlamaCpp instance if successful,
                            None otherwise.

    Raises:
        Logs errors and displays Streamlit error messages if the file is not found
        or if initialization fails.
    """
    if not os.path.exists(config.LLAMA_MODEL_PATH):
        error_msg = f"LLM model file not found at specified path: {config.LLAMA_MODEL_PATH}. Ensure the file exists and the path in config.py or the MODEL_DIRECTORY environment variable is correct."
        st.error(error_msg)
        logging.error(error_msg)
        return None
    try:
        logging.info(f"Initializing LlamaCpp from: {config.LLAMA_MODEL_PATH}")
        llm = LlamaCpp(
            streaming=True, # Enable streaming for interactive chat
            model_path=config.LLAMA_MODEL_PATH,
            temperature=config.LLAMA_TEMPERATURE,
            top_p=config.LLAMA_TOP_P,
            verbose=True, # Set to False for cleaner logs in production
            n_ctx=config.LLAMA_N_CTX,
            # Add other LlamaCpp parameters if needed (e.g., n_gpu_layers)
        )
        logging.info("LlamaCpp LLM initialized successfully.")
        return llm
    except Exception as e:
        logging.error(f"Error initializing LlamaCpp: {e}", exc_info=True)
        st.error(f"Error initializing LlamaCpp LLM: {e}. Ensure the model path is correct and the llama-cpp-python package is installed correctly (requires compilation).")
        return None

@st.cache_resource
def load_embedding_model():
    """
    Loads the HuggingFace sentence transformer embedding model.

    Uses Sentence Transformers library to download/load the specified model.
    Caches the model resource using Streamlit.

    Returns:
        Optional[HuggingFaceEmbeddings]: The initialized embedding model instance,
                                         or None if loading fails.

    Raises:
        Logs errors and displays Streamlit error messages on failure.
    """
    try:
        logging.info(f"Loading embedding model: {config.EMBEDDING_MODEL_NAME}")
        # Specify cache folder for sentence-transformers if desired
        # cache_folder = './st_embedding_cache'
        # embeddings = HuggingFaceEmbeddings(model_name=config.EMBEDDING_MODEL_NAME, cache_folder=cache_folder)
        embeddings = HuggingFaceEmbeddings(model_name=config.EMBEDDING_MODEL_NAME)
        logging.info("HuggingFace embedding model loaded successfully.")
        return embeddings
    except Exception as e:
        logging.error(f"Error loading embedding model ({config.EMBEDDING_MODEL_NAME}): {e}", exc_info=True)
        st.error(f"Error loading embedding model: {e}. Check model name and network connection.")
        return None

# --- Helper Functions ---

def generate_file_id(file):
    """
    Generates a unique MD5 hash based on the file's content.

    Reads the entire file content to compute the hash. Resets the file
    pointer afterwards so the file can be read again.

    Args:
        file (streamlit.uploaded_file_manager.UploadedFile): The uploaded file object.

    Returns:
        Optional[str]: The hex digest of the MD5 hash, or None if an error occurs.
    """
    try:
        file_content = file.getvalue() # Read file content into memory
        file.seek(0) # IMPORTANT: Reset stream position after reading
        return hashlib.md5(file_content).hexdigest()
    except Exception as e:
        logging.error(f"Error generating file ID for {getattr(file, 'name', 'Unknown File')}: {e}")
        return None

def get_vector_store_path(file_id: str) -> str:
    """
    Constructs the full path for storing/retrieving a cached vector store file.

    Args:
        file_id (str): The unique ID of the file.

    Returns:
        str: The path to the pickle file in the cache directory.
    """
    return os.path.join(config.VECTOR_STORE_CACHE_DIR, f"{file_id}.pkl")

# --- File Processing Logic ---

def transcribe_audio_file(whisper_pipe, file_path: str) -> str:
    """
    Transcribes a single audio file using the loaded Whisper pipeline.

    Args:
        whisper_pipe (transformers.Pipeline): The loaded Whisper ASR pipeline.
        file_path (str): Path to the temporary audio file to transcribe.

    Returns:
        str: The transcribed text, or an empty string if transcription fails.
    """
    if whisper_pipe is None:
        st.error("Whisper pipeline not loaded. Cannot transcribe audio.")
        logging.error("Attempted to transcribe audio, but Whisper pipeline is None.")
        return ""
    try:
        file_basename = os.path.basename(file_path)
        logging.info(f"Starting transcription for file: {file_basename}")
        # Call the pipeline with the specified language from config
        result = whisper_pipe(file_path, generate_kwargs={"language": config.WHISPER_LANGUAGE})
        transcript = result.get('text', '').strip()
        if transcript:
             logging.info(f"Transcription successful for {file_basename}. Length: {len(transcript)} chars.")
        else:
             logging.warning(f"Transcription resulted in empty text for {file_basename}.")
        return transcript
    except Exception as e:
        logging.error(f"Error transcribing file {os.path.basename(file_path)}: {e}", exc_info=True)
        st.warning(f"Could not transcribe file {os.path.basename(file_path)}. Error: {e}")
        return ""

def extract_text_from_document(file_path: str, file_type: str) -> str:
    """
    Extracts text content from various document types (PDF, DOCX, TXT).

    Uses appropriate libraries for each format and falls back to `textract`
    for other potential types (requires OS-level dependencies).

    Args:
        file_path (str): Path to the temporary document file.
        file_type (str): The MIME type of the file (e.g., 'application/pdf').

    Returns:
        str: The extracted text content, or an empty string if extraction fails.
    """
    text = ""
    file_basename = os.path.basename(file_path)
    try:
        logging.info(f"Starting text extraction from: {file_basename} (Type: {file_type})")
        if file_type == "application/pdf":
            loader = PyPDFLoader(file_path)
            documents = loader.load() # Returns list of LangChain Document objects
            text = "\n\n".join([doc.page_content for doc in documents]) # Join pages
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = DocxDocument(file_path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text])
        elif file_type == "text/plain":
            # Attempt to read with UTF-8, ignore errors for broader compatibility
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        else:
            # Fallback using textract (ensure dependencies are installed - see README)
            logging.info(f"Attempting fallback text extraction with textract for: {file_basename}")
            try:
                raw_text = textract.process(file_path)
                text = raw_text.decode('utf-8', errors='ignore')
            except Exception as te:
                logging.error(f"Textract failed for {file_basename}: {te}")
                st.warning(f"Could not extract text from {file_basename} using textract fallback. Ensure textract dependencies (like pdftotext, antiword) are installed.")
                return ""

        if text:
            logging.info(f"Text extraction successful for {file_basename}. Length: {len(text)} chars.")
        else:
            logging.warning(f"Text extraction resulted in empty text for {file_basename}.")
        return text.strip() # Return stripped text
    except Exception as e:
        logging.error(f"Error extracting text from {file_basename}: {e}", exc_info=True)
        st.warning(f"Could not extract text from file {file_basename}. Error: {e}")
        return ""

def create_or_load_vector_store(file_id: str, text_content: str, embedding_model) -> Optional[FAISS]:
    """
    Creates a new FAISS vector store or loads an existing one from the cache directory.

    If a cached store exists for the file_id, it's loaded. Otherwise, a new store
    is created using the provided text content and embedding model, then saved to cache.

    Args:
        file_id (str): The unique identifier for the file.
        text_content (str): The text extracted from the file.
        embedding_model (HuggingFaceEmbeddings): The loaded embedding model instance.

    Returns:
        Optional[FAISS]: The loaded or newly created FAISS vector store, or None if
                         creation/loading fails or text content is empty.
    """
    store_path = get_vector_store_path(file_id)
    vector_store: Optional[FAISS] = None

    # 1. Attempt to load from cache
    if os.path.exists(store_path):
        try:
            logging.debug(f"Attempting to load cached vector store from: {store_path}")
            with open(store_path, "rb") as f:
                vector_store = pickle.load(f)
            # Basic validation: Check if it's a FAISS instance and has an index
            if isinstance(vector_store, FAISS) and hasattr(vector_store, 'index') and vector_store.index is not None:
                logging.info(f"Successfully loaded cached vector store for file ID {file_id}.")
            else:
                 logging.warning(f"Cached file {store_path} for file ID {file_id} is invalid or corrupted. Recreating.")
                 vector_store = None # Force recreation
        except (pickle.UnpicklingError, EOFError, AttributeError, ImportError, IndexError) as e:
            logging.warning(f"Error loading vector store from cache {store_path}: {e}. Cache file might be corrupted or incompatible. Recreating...")
            vector_store = None # Ensure recreation if loading fails
        except Exception as e:
             logging.error(f"Unexpected error loading vector store from {store_path}: {e}. Recreating...", exc_info=True)
             vector_store = None

    # 2. If not loaded from cache, create a new store
    if vector_store is None:
        if not text_content or not text_content.strip():
            logging.warning(f"No text content provided for file ID {file_id}. Cannot create vector store.")
            return None
        if embedding_model is None:
            logging.error(f"Embedding model is not available for file ID {file_id}. Cannot create vector store.")
            st.error("Embedding model error. Cannot process file.")
            return None

        try:
            logging.info(f"Creating new vector store for file ID {file_id}.")
            # Simple approach: Treat the entire file content as one document.
            # For large files, splitting into chunks is recommended:
            # from langchain.text_splitter import RecursiveCharacterTextSplitter
            # text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
            # texts = text_splitter.split_text(text_content)
            # documents = [Document(page_content=t, metadata={"source": file_id}) for t in texts] # Add source metadata
            documents = [Document(page_content=text_content, metadata={"source": file_id})] # Add source ID as metadata

            if not documents:
                 logging.warning(f"Text content resulted in zero documents for file ID {file_id}. Cannot create vector store.")
                 return None

            # Create FAISS index from documents
            vector_store = FAISS.from_documents(documents, embedding_model)
            logging.info(f"Successfully created new vector store for file ID {file_id} with {vector_store.index.ntotal} vectors.")

            # 3. Save the newly created store to cache
            try:
                with open(store_path, "wb") as f:
                    pickle.dump(vector_store, f)
                logging.info(f"Saved newly created vector store to cache: {store_path}")
            except Exception as e_save:
                 logging.error(f"Error saving newly created vector store to {store_path}: {e_save}", exc_info=True)
                 st.warning(f"Could not cache vector store for file ID {file_id}. It will be recreated next time.")
                 # Continue with the in-memory store, but caching failed.

        except Exception as e:
            logging.error(f"Error creating FAISS vector store for file ID {file_id}: {e}", exc_info=True)
            st.error(f"Failed to create vector store for a file. Error: {e}")
            return None

    return vector_store


def process_uploaded_files(uploaded_files: List[st.runtime.uploaded_file_manager.UploadedFile],
                           file_type_handler, # Function to call (transcribe or extract)
                           embedding_model,
                           **kwargs) -> Tuple[List[str], List[Dict]]:
    """
    Processes a list of uploaded files.

    For each file:
    1. Generates a unique ID based on content.
    2. Skips if already active in the session.
    3. Tries to load its vector store from cache.
    4. If not cached, uses a temporary file and `file_type_handler` to get text.
    5. Creates a new vector store from the text and caches it.
    6. Stores the loaded/created vector store in session state.

    Args:
        uploaded_files: List of Streamlit UploadedFile objects.
        file_type_handler: The function to call for text extraction/transcription
                           (e.g., `transcribe_audio_file` or `extract_text_from_document`).
        embedding_model: The loaded HuggingFace embedding model.
        **kwargs: Additional arguments to pass to the `file_type_handler` (like whisper_pipe).

    Returns:
        Tuple containing:
            - List[str]: A list of file IDs successfully processed or loaded in this batch.
            - List[Dict]: A summary report of processing status for each file.
    """
    processed_file_ids_in_batch = []
    processing_summary = []

    if not uploaded_files:
        return processed_file_ids_in_batch, processing_summary

    if embedding_model is None:
        st.error("Embedding model not loaded. Cannot process files.")
        logging.error("Attempted file processing, but embedding model is None.")
        return processed_file_ids_in_batch, processing_summary

    for uploaded_file in uploaded_files:
        file_name = uploaded_file.name
        file_id = generate_file_id(uploaded_file)

        if not file_id:
            st.error(f"Could not generate unique ID for file: {file_name}")
            processing_summary.append({"name": file_name, "id": "N/A", "status": "Error: ID generation failed"})
            continue

        # Store mapping from ID to original filename for display purposes
        st.session_state.file_id_to_name[file_id] = file_name

        # Check if this file ID is already marked as active in the session
        if file_id in st.session_state.get('active_file_ids', set()):
             logging.info(f"File '{file_name}' ({file_id}) is already active in the session. Skipping reprocessing.")
             processing_summary.append({"name": file_name, "id": file_id, "status": "Already Active"})
             # Ensure it's included in the batch list if needed for merging logic later
             if file_id not in processed_file_ids_in_batch:
                 processed_file_ids_in_batch.append(file_id)
             # Make sure its vector store is loaded into the session state if not already
             if file_id not in st.session_state.active_vector_stores:
                 store = create_or_load_vector_store(file_id, "", embedding_model) # Try loading from cache
                 if store:
                      st.session_state.active_vector_stores[file_id] = store
                 else:
                      logging.warning(f"Could not load vector store for already active file ID {file_id}. It might be missing from cache.")
                      # Attempt to remove it from active_file_ids if store is unavailable?
                      # This could happen if cache was deleted manually.
                      # st.session_state.active_file_ids.discard(file_id)
             continue # Skip to the next file

        # --- Processing logic for new files ---
        vector_store: Optional[FAISS] = None
        status = "Pending"
        text_content = ""

        # 1. Try loading from cache first (create_or_load handles this internally)
        # We pass empty text initially; if cache exists, it loads; if not, it returns None.
        cached_store = create_or_load_vector_store(file_id, "", embedding_model)

        if cached_store:
            logging.info(f"Loaded cached vector store for new file: {file_name} ({file_id}).")
            vector_store = cached_store
            status = "Loaded from Cache"
        else:
            # 2. If not cached, process the file to get text
            logging.info(f"Processing new file (not found in cache): {file_name} ({file_id}).")
            try:
                # Use a temporary file to pass to the handler functions
                # This avoids potential issues with different library expectations (path vs buffer)
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file_name)[1]) as temp_file:
                    temp_file.write(uploaded_file.getvalue())
                    temp_file_path = temp_file.name

                # Call the appropriate handler (transcribe or extract text)
                # Pass necessary kwargs (like whisper_pipe for audio)
                text_content = file_type_handler(temp_file_path=temp_file_path, file_type=uploaded_file.type, **kwargs)

                os.remove(temp_file_path) # Clean up the temporary file

                if not text_content:
                    status = "Warning: No text extracted"
                    logging.warning(f"File '{file_name}' resulted in no text content after processing.")
                    # Don't try to create a vector store if no text

            except Exception as e:
                logging.error(f"Error during file processing ({file_name}): {e}", exc_info=True)
                st.error(f"Failed to process file: {file_name}. Error: {e}")
                status = f"Error: Processing failed ({e})"
                # Ensure we continue to the next file on error
                processing_summary.append({"name": file_name, "id": file_id, "status": status})
                continue

            # 3. Create and cache the vector store for the newly processed file (if text exists)
            if text_content:
                # Call create_or_load again, this time with text_content to force creation
                vector_store = create_or_load_vector_store(file_id, text_content, embedding_model)
                if vector_store:
                    status = "Processed and Cached"
                else:
                    status = "Error: Vector store creation failed"
                    logging.error(f"Failed to create vector store for file {file_name} even after processing.")
            else:
                 # status remains "Warning: No text extracted"
                 pass # No vector store to create

        # 4. Store successfully loaded/created vector store in session state
        if vector_store:
            st.session_state.active_vector_stores[file_id] = vector_store
            processed_file_ids_in_batch.append(file_id)
        else:
            # Ensure file ID isn't associated with a failed store creation/loading
            if file_id in st.session_state.active_vector_stores:
                 del st.session_state.active_vector_stores[file_id]
            # Ensure status reflects the failure if it wasn't already set
            if status.startswith("Pending") or status.startswith("Loaded"):
                 status = "Error: Failed to load/create vector store"


        processing_summary.append({"name": file_name, "id": file_id, "status": status})
        logging.debug(f"File: {file_name}, ID: {file_id}, Status: {status}")

    return processed_file_ids_in_batch, processing_summary


def merge_active_vector_stores(embedding_model) -> Optional[FAISS]:
    """
    Merges multiple individual FAISS vector stores held in session state into a single store.

    Retrieves all vector stores from `st.session_state.active_vector_stores`.
    If multiple valid stores exist, it initializes a new FAISS index and uses the
    `merge_from` method to combine vectors and document mappings. This enables
    the RAG chain to retrieve context across all active files.

    Args:
        embedding_model: The loaded HuggingFace embedding model instance. Needed for
                         potential re-initialization or validation.

    Returns:
        Optional[FAISS]: The merged FAISS vector store if successful and non-empty,
                         or a single store if only one was active, or None if no
                         stores are active or merging fails.
    """
    active_stores: Dict[str, FAISS] = st.session_state.get('active_vector_stores', {})
    valid_stores_to_merge = {
        fid: store for fid, store in active_stores.items()
        if isinstance(store, FAISS) and hasattr(store, 'index') and store.index is not None and store.index.ntotal > 0
    }

    num_valid_stores = len(valid_stores_to_merge)
    logging.info(f"Found {num_valid_stores} valid active vector stores to potentially merge.")

    if num_valid_stores == 0:
        logging.info("No active vector stores with content found to merge.")
        return None
    if embedding_model is None:
         st.error("Embedding model not loaded. Cannot merge vector stores.")
         logging.error("Attempted merging vector stores, but embedding model is None.")
         return None

    if num_valid_stores == 1:
        file_id, single_store = list(valid_stores_to_merge.items())[0]
        logging.info(f"Only one active vector store ({st.session_state.file_id_to_name.get(file_id, file_id)}). No merge needed.")
        return single_store

    # --- Merge multiple stores ---
    logging.info(f"Merging {num_valid_stores} active vector stores...")
    try:
        # Get the first valid store to potentially initialize the merged store structure
        first_file_id, first_store = list(valid_stores_to_merge.items())[0]
        stores_to_merge_iter = iter(valid_stores_to_merge.items())
        next(stores_to_merge_iter) # Consume the first item

        # Initialize the merged store by copying the structure of the first one
        # This seems more robust than creating from scratch if `merge_from` expects it
        merged_store = FAISS(
            embedding_function=first_store.embedding_function,
            index=first_store.index,
            docstore=first_store.docstore,
            index_to_docstore_id=first_store.index_to_docstore_id
        )
        logging.debug(f"Initialized merged store using structure from file: {st.session_state.file_id_to_name.get(first_file_id, first_file_id)}")


        # Merge the remaining stores into the initialized one
        for file_id, store_to_merge in stores_to_merge_iter:
            file_name = st.session_state.file_id_to_name.get(file_id, file_id)
            logging.debug(f"Merging store from file: {file_name} (ntotal={store_to_merge.index.ntotal})")
            try:
                merged_store.merge_from(store_to_merge)
            except Exception as merge_err:
                logging.error(f"Failed to merge store from file '{file_name}': {merge_err}", exc_info=True)
                st.warning(f"Could not merge data from file: {file_name}. Skipping.")
                # Continue merging other stores

        final_vector_count = merged_store.index.ntotal
        logging.info(f"Merging complete. Final merged store contains {final_vector_count} vectors.")
        if final_vector_count == 0:
             logging.warning("Merged store is empty after merging process (vectors might be zero or merge failed).")
             st.warning("The combined context appears empty after processing. Check file content or processing logs.")
             return None # Return None if the merged store ended up empty

        return merged_store

    except Exception as e:
        logging.error(f"Error during vector store merge process: {e}", exc_info=True)
        st.error(f"An error occurred while combining file contexts: {e}")
        return None

# --- LangChain Setup ---

def get_conversational_chain(llm: LlamaCpp, vector_store: FAISS) -> Optional[ConversationalRetrievalChain]:
    """
    Initializes the LangChain ConversationalRetrievalChain.

    Uses the provided LLM and the (potentially merged) FAISS vector store as a retriever.
    Sets up conversation memory to maintain context across turns.

    Args:
        llm: The initialized LlamaCpp model instance.
        vector_store: The FAISS vector store containing the combined context.

    Returns:
        Optional[ConversationalRetrievalChain]: The initialized chain object,
                                                or None if initialization fails.
    """
    if vector_store is None:
        st.error("Combined vector store is not available. Cannot initialize conversation chain.")
        logging.warning("Attempted to initialize chain with vector_store=None.")
        return None
    if llm is None:
        st.error("LLM is not loaded. Cannot initialize conversation chain.")
        logging.warning("Attempted to initialize chain with llm=None.")
        return None

    try:
        logging.info("Initializing ConversationalRetrievalChain...")
        # Create a new memory buffer for each chain instance.
        # Could potentially be stored in session state for longer persistence if needed.
        memory = ConversationBufferMemory(
            memory_key="chat_history", # Key LangChain uses to store history
            return_messages=True,      # Return history as Message objects
            output_key='answer'        # Specify the key for the LLM's answer
        )

        # Create the retriever from the vector store
        retriever = vector_store.as_retriever(
            search_kwargs={"k": config.RETRIEVER_K} # Use K from config
        )

        # Assemble the chain
        chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=retriever,
            memory=memory,
            return_source_documents=True, # Set to True to see retrieved chunks (good for debugging)
            verbose=False # Set to True for very detailed LangChain logs
        )
        logging.info("ConversationalRetrievalChain initialized successfully.")
        return chain
    except Exception as e:
        logging.error(f"Error initializing ConversationalRetrievalChain: {e}", exc_info=True)
        st.error(f"Could not initialize the conversation engine: {e}")
        return None

# --- Chat Logic ---

def handle_conversation(query: str, chain: ConversationalRetrievalChain, history: List[Tuple[str, str]]):
    """
    Processes a user query using the conversational chain and updates history.

    Args:
        query (str): The user's input question.
        chain (ConversationalRetrievalChain): The active LangChain conversation chain.
        history (List[Tuple[str, str]]): The display history (list of (user, bot) tuples).
                                         Note: Chain manages its internal memory separately.

    Returns:
        str: The bot's generated response.
    """
    if not query:
        return "Please enter a question."
    if not chain:
        # Check if the underlying vector store exists but chain creation failed
        if st.session_state.get('combined_vector_store'):
             error_msg = "Chat engine initialization failed. Please check application logs or model configuration."
             logging.warning(error_msg)
             return error_msg
        else:
             # This case means no files were processed successfully
             error_msg = "Chat engine is not ready. Please upload and process files successfully first."
             logging.warning(error_msg)
             return error_msg

    try:
        logging.info(f"Running chain with query: '{query}'")
        # The chain uses its internal memory for conversation history.
        # We pass the new question here.
        result = chain({"question": query})

        answer = result.get("answer", "Sorry, I couldn't generate an answer for that.")
        logging.info(f"Chain returned answer snippet: '{answer[:100]}...'")

        # Append interaction to the display history list (managed in session state)
        history.append((query, answer))

        # Optional: Log source documents for debugging relevance
        if 'source_documents' in result and result['source_documents']:
             logging.debug(f"Retrieved {len(result['source_documents'])} source documents for query: '{query}'")
             for i, doc in enumerate(result['source_documents']):
                 source_id = doc.metadata.get('source', 'Unknown Source')
                 source_name = st.session_state.file_id_to_name.get(source_id, source_id)
                 logging.debug(f"  Source Doc {i+1} [File: {source_name}]: {doc.page_content[:150]}...")
        else:
             logging.debug(f"No source documents were retrieved for query: '{query}'")


        return answer
    except Exception as e:
        logging.error(f"Error during conversation chain execution: {e}", exc_info=True)
        st.error(f"An error occurred while getting the response: {e}")
        return "Sorry, an error occurred while processing your request. Please check the logs."

# --- Streamlit UI Components ---

def initialize_session_state():
    """
    Initializes Streamlit session state variables on first run or if they get cleared.
    Uses descriptive keys for clarity.
    """
    defaults = {
        'chat_display_history': [],        # Stores (user_query, bot_response) tuples for display
        'generated_responses': ["Hello! Upload audio or documents, then ask me questions about their combined content."], # Bot responses for streamlit-chat
        'past_user_inputs': ["Hi!"],       # User inputs for streamlit-chat
        'active_file_ids': set(),          # Stores file_ids (hashes) of currently active files
        'active_vector_stores': {},        # Maps file_id -> loaded FAISS object for individual files
        'file_id_to_name': {},             # Maps file_id -> original filename for display
        'combined_vector_store': None,     # The merged FAISS store for the current active files
        'conversation_chain': None         # The active LangChain ConversationalRetrievalChain instance
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            # Use deep copies for mutable defaults if necessary, though sets/dicts handled ok here
            st.session_state[key] = value
    logging.debug("Session state initialized or verified.")


def display_chat_ui():
    """Displays the chat interface using streamlit-chat and handles user input submission."""
    st.markdown("### Chat Interface")
    st.write("Ask questions based on the content of the currently active files listed in the sidebar.")
    st.write("---")

    # Containers for chat history and user input
    response_container = st.container() # Where bot responses will appear
    input_container = st.container()    # Where user types input

    with input_container:
        # Use a form for better control over input submission
        with st.form(key='chat_input_form', clear_on_submit=True):
            user_input = st.text_input("Your Question:", placeholder="e.g., 'Summarize the main points from the audio file.'", key='user_input')
            submit_button = st.form_submit_button(label='Send')

        if submit_button and user_input:
            # Retrieve the current chain from session state
            current_chain = st.session_state.get('conversation_chain')
            if current_chain:
                with st.spinner('Thinking... Please wait.'):
                    # Pass the display history list directly if handle_conversation needs it
                    # Although chain manages internal memory, display history is separate
                    output = handle_conversation(user_input, current_chain, st.session_state['chat_display_history'])

                # Update the lists used by streamlit-chat for display
                st.session_state['past_user_inputs'].append(user_input)
                st.session_state['generated_responses'].append(output)
                print(f"User: {user_input} | Bot: {output}") # Debug print for local testing

                # Rerun the script immediately to update the chat display
                st.rerun()
            else:
                 # Provide feedback if the chain isn't ready
                 st.warning("The chat engine isn't ready. Please ensure files have been uploaded and processed successfully.")

    # Display the chat history using streamlit-chat messages
    if st.session_state['generated_responses']:
        with response_container:
            # Loop through the generated responses to display the chat history
            for i in range(len(st.session_state['generated_responses'])):
                # Display user message first, if available
                if i < len(st.session_state['past_user_inputs']):
                    message(st.session_state["past_user_inputs"][i],
                            is_user=True,
                            key=f"{i}_user",
                            avatar_style="thumbs") # Or "initials", "identicon", etc.
                # Display corresponding bot response
                message(st.session_state["generated_responses"][i],
                        key=str(i),
                        avatar_style="fun-emoji") # Or "bottts", "croodles", etc.


def clear_chat_session():
    """
    Resets the session state related to files and conversation history.
    Keeps loaded models in cache.
    """
    logging.info("Clearing chat session state (files, history, chain).")
    # Reset file tracking and vector stores
    st.session_state.active_file_ids = set()
    st.session_state.active_vector_stores = {}
    st.session_state.file_id_to_name = {}
    st.session_state.combined_vector_store = None
    st.session_state.conversation_chain = None

    # Reset chat display history
    st.session_state.chat_display_history = []
    st.session_state.generated_responses = ["Session cleared. Upload new files to begin."]
    st.session_state.past_user_inputs = ["Cleared!"]

    # Clear Streamlit's file uploader state requires a rerun trick or specific keys
    # This is often handled implicitly by rerun, but good to be aware
    st.cache_data.clear() # Clear function caches if needed, though less relevant here
    # st.cache_resource.clear() # Avoid clearing models unless intended

    logging.info("Session state cleared. Triggering rerun.")
    # Rerun the app to reflect the cleared state in the UI
    st.rerun()


# --- Main Application Flow ---

def main():
    # --- Page Configuration (do this first) ---
    st.set_page_config(
        page_title="Chat with Audio & Docs",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    st.title("📊 Chat with Multiple Audio & Document Files")
    st.markdown("""
        Upload audio files (like MP3, WAV) and/or documents (PDF, DOCX, TXT).
        The application will process them, and you can then ask questions based on the **combined content**
        of all files currently active in the session. Processing may take time depending on file size and type.
        Use the sidebar to manage files and see processing status.
    """)
    st.divider()

    # --- Initialization ---
    # Ensure session state variables exist
    initialize_session_state()

    # --- Load Models (Cached Resources) ---
    # These functions are cached, so they only run computation once.
    # Errors during loading are handled internally and displayed.
    with st.spinner("Loading AI models... This might take a moment on first run."):
        whisper_pipeline = load_whisper_pipeline()
        llm = load_llama_cpp()
        embedding_model = load_embedding_model()

    # --- Sidebar UI Definition ---
    with st.sidebar:
        st.header("⚙️ Upload & Manage Files")

        # File Uploaders
        uploaded_audio_files = st.file_uploader(
            "1. Upload Audio Files",
            accept_multiple_files=True,
            type=['mp3', 'wav', 'mp4', 'm4a', 'ogg'], # Added ogg
            key='audio_uploader_key', # Unique key helps Streamlit manage state
            help="Upload one or more audio files for transcription."
        )
        uploaded_document_files = st.file_uploader(
            "2. Upload Document Files",
            accept_multiple_files=True,
            type=['pdf', 'docx', 'txt'],
            key='doc_uploader_key', # Unique key
            help="Upload one or more documents for text extraction."
        )

        st.divider()

        # Session Control
        st.header("🔄 Session Control")
        if st.button("Clear Session & Start Over", key="clear_session_button", help="Removes all uploaded files and chat history from the current session."):
            clear_chat_session()

        st.divider()

        # Active Files Display
        st.header("📁 Active Files in Context")
        if not st.session_state.active_file_ids:
            st.info("No files are currently active. Upload files to begin.")
        else:
            st.markdown("**Currently using content from:**")
            # Display names of active files
            for file_id in sorted(list(st.session_state.active_file_ids)): # Sort for consistent order
                file_name = st.session_state.file_id_to_name.get(file_id, f"ID: {file_id[:8]}...")
                st.write(f" - `{file_name}`")
            st.caption(f"Total: {len(st.session_state.active_file_ids)} file(s)")

        # Placeholder for processing summary
        st.divider()
        st.header("📊 Processing Status")
        processing_status_placeholder = st.empty()
        processing_status_placeholder.info("Upload files to see processing status here.")

    # --- File Processing and Context Update Logic ---
    # This section runs *after* the UI elements are defined.
    # It checks if new files were uploaded in this run.

    newly_processed_file_ids = set() # Track IDs successfully processed *in this specific run*
    processing_summary_current_run = [] # Summary for *this run's* uploads

    # Determine if new files were uploaded in this interaction
    process_audio = uploaded_audio_files is not None and len(uploaded_audio_files) > 0
    process_docs = uploaded_document_files is not None and len(uploaded_document_files) > 0

    files_were_uploaded = process_audio or process_docs

    if files_were_uploaded:
        processing_status_placeholder.info("Processing uploaded files... Please wait.")
        with st.spinner("Processing files... This can take some time."):
            # Process Audio Files if any were uploaded
            if process_audio:
                logging.info(f"Processing {len(uploaded_audio_files)} uploaded audio file(s).")
                # Define the handler function for audio processing
                def audio_handler(temp_file_path, file_type, **kwargs): # Match expected args
                    # Ensure whisper_pipeline is passed correctly via kwargs
                    pipe = kwargs.get('whisper_pipe')
                    if pipe is None:
                        logging.error("Whisper pipeline missing in audio_handler kwargs.")
                        return ""
                    return transcribe_audio_file(pipe, temp_file_path)

                processed_ids, summary = process_uploaded_files(
                    uploaded_audio_files,
                    audio_handler,
                    embedding_model,
                    whisper_pipe=whisper_pipeline # Pass the loaded pipeline here
                )
                newly_processed_file_ids.update(processed_ids)
                processing_summary_current_run.extend(summary)

            # Process Document Files if any were uploaded
            if process_docs:
                logging.info(f"Processing {len(uploaded_document_files)} uploaded document file(s).")
                # Define the handler function for document processing
                def document_handler(temp_file_path, file_type, **kwargs): # Match expected args
                     # No extra kwargs needed here currently
                    return extract_text_from_document(temp_file_path, file_type)

                processed_ids, summary = process_uploaded_files(
                    uploaded_document_files,
                    document_handler,
                    embedding_model
                    # No extra kwargs needed here currently
                )
                newly_processed_file_ids.update(processed_ids)
                processing_summary_current_run.extend(summary)

        # --- Update Active Files & Rebuild Context if needed ---
        # Check if any *new* files were successfully processed OR if the set of active files needs updating
        # (e.g., if a previously cached file was uploaded again and successfully loaded)
        active_files_changed = False
        if newly_processed_file_ids:
             # Add newly processed valid IDs to the main active set in session state
            current_active_set = st.session_state.get('active_file_ids', set())
            updated_active_set = current_active_set.union(newly_processed_file_ids)

            # Check if the active set actually changed compared to before this run
            if updated_active_set != current_active_set:
                st.session_state.active_file_ids = updated_active_set
                active_files_changed = True
                logging.info(f"Active file set changed. New set: {st.session_state.active_file_ids}")
            else:
                 logging.info("Newly processed files were already active. No change to active set.")

        # If the set of active files changed, we MUST rebuild the merged context and chain
        if active_files_changed:
            logging.info("Active files changed, rebuilding combined context and conversation chain.")
            with st.spinner("Combining context from active files..."):
                # Merge vector stores based on the *updated* active_file_ids and loaded stores
                st.session_state.combined_vector_store = merge_active_vector_stores(embedding_model)

            if st.session_state.combined_vector_store:
                 with st.spinner("Initializing conversation engine..."):
                    # Re-initialize the chain with the new combined store
                    st.session_state.conversation_chain = get_conversational_chain(llm, st.session_state.combined_vector_store)
                 if st.session_state.conversation_chain:
                     st.success("Chat engine ready with updated file context!")
                 else:
                     # Merging succeeded, but chain init failed
                     st.error("Failed to initialize chat engine after updating context. Check LLM/LangChain setup.")
                     # Keep the merged store, but clear the broken chain
                     st.session_state.conversation_chain = None
            else:
                 # Merging failed or resulted in an empty store
                 st.warning("Could not create a combined context from the active files. Chat may not work.")
                 # Ensure chain is cleared if merging failed
                 st.session_state.conversation_chain = None

            # Update the processing status in the sidebar *after* processing is complete
            if processing_summary_current_run:
                 status_text = "##### Last Upload Summary:\n"
                 for info in processing_summary_current_run:
                     status_icon = "✅" if "Processed" in info['status'] or "Loaded" in info['status'] or "Active" in info['status'] else \
                                   "⚠️" if "Warning" in info['status'] else "❌"
                     status_text += f"- {status_icon} `{info['name']}`: {info['status']}\n"
                 processing_status_placeholder.markdown(status_text)
            else:
                 # This case shouldn't happen if files were uploaded, but as a fallback:
                 processing_status_placeholder.warning("Processing finished, but no summary generated.")


            # Rerun ONLY IF the active files actually changed to update sidebar & potentially chat readiness
            st.rerun()

        else:
             # Files were uploaded, but they were already active or processing failed entirely.
             # Update status but don't trigger full context rebuild/rerun unless necessary.
             if processing_summary_current_run:
                  status_text = "##### Last Upload Summary:\n"
                  for info in processing_summary_current_run:
                      status_icon = "✅" if "Active" in info['status'] else \
                                    "⚠️" if "Warning" in info['status'] else "❌"
                      status_text += f"- {status_icon} `{info['name']}`: {info['status']}\n"
                  processing_status_placeholder.markdown(status_text)
             elif files_were_uploaded: # Files uploaded but summary is empty (e.g., all failed ID gen)
                 processing_status_placeholder.error("Processing attempted, but encountered errors. Check logs.")
             # No rerun here, let the chat UI display normally

    # --- Display Chat UI ---
    # This part runs regardless of file uploads in the current run,
    # using the latest chain available in the session state.
    display_chat_ui()


if __name__ == "__main__":
    # Ensure potential environment variables from a .env file are loaded *before* config is used heavily
    # Optional: uncomment if using python-dotenv and a .env file
    # from dotenv import load_dotenv
    # load_dotenv()

    main()